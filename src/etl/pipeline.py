# REQ_F001: ETL pipeline — Extract, Transform, Load for local documents
# REQ_F002: Chunking and embedding for vector search

import shutil
import os
from pathlib import Path
from dotenv import load_dotenv  #read .env config file
from src.metadata.repository import load_document_metadata

from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
#TextLoader read .txt files into LangChain Document objects
#PyPDFLoader read .pdf files into LangChain Document objects
#Docx allow use of python-docx to extract text then wrap into LangChain Document objects

from langchain_text_splitters import RecursiveCharacterTextSplitter
#split long documents into smaller overlapping chunks

from src.vector.factory import get_vector_backend
# ── AZURE SWAP ──
# Replace HuggingFaceEmbeddings with:
#   from langchain_openai import AzureOpenAIEmbeddings
# Requires: AZURE_OPENAI_API_KEY, AZURE_OPENAI_ENDPOINT in .env
# ── END AZURE SWAP ──


#load all values from .env into environment variables
load_dotenv() 


def get_active_metadata_filenames() -> set[str]:
    """Return filenames that belong to active metadata records only."""
    active_documents = load_document_metadata()

    return {
        document["filename"]
        for document in active_documents
    }


def extract_docx_text(file_path: Path) -> str:
    """Extract searchable paragraph and table text from a DOCX file."""
    docx_file = DocxDocument(str(file_path))
    text_parts = []

    for paragraph in docx_file.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            text_parts.append(paragraph_text)

    for table in docx_file.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )

            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)

#function to handle E in ETL
def load_documents(
    folder_path: str,
    allowed_filenames: set[str] | None = None,
) -> tuple[list, int]:
    #(REQ_F001) E in ETL, read all .txt and .pdf files from given folder
    #Return a list of LangChain Document object, each holding file text + metadata
    
    #list for storing all loaded Document objects
    documents = []

    #count physical files separately because PDFs can produce one Document per page
    source_file_count = 0

    #converet string path to Path object 
    folder = Path(folder_path)

    #loop through every file in the folder
    for file_path in folder.iterdir():
        if not file_path.is_file():
          continue

        if allowed_filenames is not None and file_path.name not in allowed_filenames:
            continue

        suffix = file_path.suffix.lower()

        if suffix not in [".txt", ".pdf", ".docx"]:
            continue

        source_file_count += 1

        if suffix == ".txt":
            loader = TextLoader(str(file_path), encoding="utf-8")
            documents.extend(loader.load())

        elif suffix == ".pdf":
            loader = PyPDFLoader(str(file_path))
            documents.extend(loader.load())

        elif suffix == ".docx":
            docx_text = extract_docx_text(file_path)

            if docx_text.strip():
                documents.append(
                    Document(
                        page_content=docx_text,
                        metadata={"source": str(file_path)},
                    )
                )
    
    #feedback for the progress
    print(
        f"Loaded {source_file_count} file(s) and "
        f"{len(documents)} document object(s) from {folder_path}"
    )
    return documents, source_file_count


def load_single_document(file_path: Path) -> list:
    """Load one supported source file into LangChain Document object(s)."""
    suffix = file_path.suffix.lower()

    if suffix == ".txt":
        loader = TextLoader(str(file_path), encoding="utf-8")
        return loader.load()

    if suffix == ".pdf":
        loader = PyPDFLoader(str(file_path))
        return loader.load()

    if suffix == ".docx":
        docx_text = extract_docx_text(file_path)

        if docx_text.strip():
            return [
                Document(
                    page_content=docx_text,
                    metadata={"source": str(file_path)},
                )
            ]

    return []


#function to handle T in ETL
def chunk_documents(documents: list) -> list:
    #(REQ_F002): splits large documents into smaller overlapping chunks.
    #smaller chunks means more precise retrieval. 
    #overlap means no lost context at boundaries
    
    splitter = RecursiveCharacterTextSplitter(
        #max characters per chunk
        chunk_size=900,

        #overlap between consecutive chunks to preserve context
        chunk_overlap=250,

        #tries to split at paragraph first, then lines, then words
        separators=["\n\n","\n"," ",""]

    )

    #new list of smaller Document objects
    chunks = splitter.split_documents(documents)
    print(f"Split into {len(chunks)} chunk(s)")
    return chunks


def embed_and_store(
    chunks: list,
    db_path: str,
    collection_name: str,
    vector_backend=None,
) -> None:
    """Persist chunked documents through the selected vector backend."""
    vector_backend = vector_backend or get_vector_backend()

    vector_backend.store_chunks(
        chunks=chunks,
        db_path=db_path,
        collection_name=collection_name,
    )


def delete_vectors_for_source(
    source_path: str,
    db_path: str | None = None,
    collection_name: str | None = None,
) -> int:
    """Delete existing vectors/index records that came from one source file."""
    db_path = db_path or os.getenv("CHROMA_DB_PATH")
    collection_name = collection_name or os.getenv("CHROMA_COLLECTION_NAME")

    if db_path is None or collection_name is None:
        raise ValueError("CHROMA_DB_PATH and CHROMA_COLLECTION_NAME are required.")

    vector_backend = get_vector_backend()

    return vector_backend.delete_vectors_for_source(
        source_path=source_path,
        db_path=db_path,
        collection_name=collection_name,
    )


def index_single_document(
    file_path: str | Path,
    db_path: str | None = None,
    collection_name: str | None = None,
) -> dict:
    """Chunk, embed, and append one source file into the configured vector backend."""
    db_path = db_path or os.getenv("CHROMA_DB_PATH")
    collection_name = collection_name or os.getenv("CHROMA_COLLECTION_NAME")

    if db_path is None or collection_name is None:
        raise ValueError("CHROMA_DB_PATH and CHROMA_COLLECTION_NAME are required.")

    source_file = Path(file_path)
    documents = load_single_document(source_file)
    chunks = chunk_documents(documents)

    if not chunks:
        return {
            "source": str(source_file),
            "document_objects_loaded": 0,
            "chunks_indexed": 0,
        }

    vector_backend = get_vector_backend()

    vector_backend.add_chunks(
        chunks=chunks,
        db_path=db_path,
        collection_name=collection_name,
    )

    return {
        "source": str(source_file),
        "document_objects_loaded": len(documents),
        "chunks_indexed": len(chunks),
    }


def index_changed_documents(source_paths: list[str]) -> dict:
    """Delete and re-index vectors for one or more changed source files."""
    unique_source_paths = list(dict.fromkeys(source_paths))

    if not unique_source_paths:
        raise ValueError("At least one source path is required.")

    update_results = []
    total_deleted_vectors = 0
    total_chunks_indexed = 0
    total_document_objects_loaded = 0

    for source_path in unique_source_paths:
        deleted_vector_count = delete_vectors_for_source(source_path)
        index_result = index_single_document(source_path)

        total_deleted_vectors += deleted_vector_count
        total_chunks_indexed += index_result["chunks_indexed"]
        total_document_objects_loaded += index_result["document_objects_loaded"]

        update_results.append(
            {
                "source": source_path,
                "deleted_vector_count": deleted_vector_count,
                "document_objects_loaded": index_result["document_objects_loaded"],
                "chunks_indexed": index_result["chunks_indexed"],
            }
        )

    return {
        "changed_document_count": len(unique_source_paths),
        "updated_sources": unique_source_paths,
        "update_results": update_results,
        "total_deleted_vectors": total_deleted_vectors,
        "total_document_objects_loaded": total_document_objects_loaded,
        "total_chunks_indexed": total_chunks_indexed,
    }


def index_changed_documents_with_cleanup(
    index_source_paths: list[str],
    cleanup_source_paths: list[str],
) -> dict:
    """Delete old/current vectors, then index only active changed source files."""
    unique_index_paths = list(dict.fromkeys(index_source_paths))
    unique_cleanup_paths = list(dict.fromkeys(cleanup_source_paths))

    if not unique_index_paths:
        raise ValueError("At least one source path is required for indexing.")

    update_results = []
    cleanup_results = []
    total_deleted_vectors = 0
    total_chunks_indexed = 0
    total_document_objects_loaded = 0

    for source_path in unique_cleanup_paths:
        deleted_vector_count = delete_vectors_for_source(source_path)
        total_deleted_vectors += deleted_vector_count

        cleanup_results.append(
            {
                "source": source_path,
                "deleted_vector_count": deleted_vector_count,
            }
        )

    for source_path in unique_index_paths:
        index_result = index_single_document(source_path)

        total_chunks_indexed += index_result["chunks_indexed"]
        total_document_objects_loaded += index_result["document_objects_loaded"]

        update_results.append(
            {
                "source": source_path,
                "deleted_vector_count": next(
                    (
                        cleanup_result["deleted_vector_count"]
                        for cleanup_result in cleanup_results
                        if cleanup_result["source"] == source_path
                    ),
                    0,
                ),
                "document_objects_loaded": index_result["document_objects_loaded"],
                "chunks_indexed": index_result["chunks_indexed"],
            }
        )

    return {
        "changed_document_count": len(unique_index_paths),
        "updated_sources": unique_index_paths,
        "cleanup_sources": unique_cleanup_paths,
        "cleanup_results": cleanup_results,
        "update_results": update_results,
        "total_deleted_vectors": total_deleted_vectors,
        "total_document_objects_loaded": total_document_objects_loaded,
        "total_chunks_indexed": total_chunks_indexed,
    }


def rebuild_vector_store(
    docs_path: str | None = None,
    db_path: str | None = None,
    collection_name: str | None = None,
    vector_backend=None,
) -> dict:
    """Rebuild the selected search index from active local working documents."""
    docs_path = docs_path or os.getenv("DOCUMENTS_PATH")
    db_path = db_path or os.getenv("CHROMA_DB_PATH")
    collection_name = collection_name or os.getenv("CHROMA_COLLECTION_NAME")

    if docs_path is None or db_path is None or collection_name is None:
        raise ValueError("DOCUMENTS_PATH and vector backend index settings are required.")

    vector_backend = vector_backend or get_vector_backend()
    vector_backend.reset_index(db_path, collection_name)

    active_filenames = get_active_metadata_filenames()
    documents, source_file_count = load_documents(docs_path, active_filenames)
    chunks = chunk_documents(documents)

    embed_and_store(
        chunks,
        db_path,
        collection_name,
        vector_backend=vector_backend,
    )

    return {
        "documents_indexed": source_file_count,
        "document_objects_loaded": len(documents),
        "chunks_indexed": len(chunks),
        "collection_name": collection_name,
        "db_path": db_path,
    }


#main entry point
if __name__ =="__main__":
    #orchestrates full ETL pipeline when run directly
    # Reads config from .env, processes all active documents, and stores them in the configured vector backend.

    #read config from .env
    docs_path = os.getenv("DOCUMENTS_PATH")
    db_path = os.getenv("CHROMA_DB_PATH")
    collection_name = os.getenv("CHROMA_COLLECTION_NAME")

    print("Starting ETL pipeline...")

    #step 1: Extract (E)
    active_filenames = get_active_metadata_filenames()
    documents, source_file_count = load_documents(docs_path, active_filenames)

    #step 2: Transform (T)
    chunks = chunk_documents(documents)

    #step 3: Load
    embed_and_store(chunks, db_path, collection_name)

    print("ETL pipeline complete. Search index is ready")
