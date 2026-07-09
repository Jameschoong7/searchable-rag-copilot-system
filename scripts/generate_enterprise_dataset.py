"""Generate a synthetic enterprise knowledge-base corpus for connector testing.

The generated files are intentionally organized by department folder/section so
OneDrive and OneNote connector ingestion can infer default ACL metadata from
paths such as `/Enterprise Knowledge Base/IT/...`.
"""

from __future__ import annotations

import csv
import re
import textwrap
from pathlib import Path

from docx import Document


ROOT = Path("datasets/enterprise_knowledge_base")
ONEDRIVE_ROOT = ROOT / "onedrive" / "Enterprise Knowledge Base"
ONENOTE_ROOT = ROOT / "onenote" / "Enterprise Knowledge Base"


DEPARTMENTS = {
    "IT": [
        (
            "Password And Account Policy",
            "Policy",
            [
                "Minimum password length is 12 characters.",
                "Passwords expire every 90 days, with reminder emails sent 14 days before expiry.",
                "Employees must not reuse any of their last 10 passwords.",
                "MFA is required for VPN, Azure portal, payroll, and knowledge-base administration.",
                "Temporary passwords must be changed during first sign-in.",
            ],
            "docx",
        ),
        (
            "VPN Access Runbook",
            "Runbook",
            [
                "VPN requests require manager approval and IT ticket reference.",
                "Standard VPN access expires after 180 days unless renewed.",
                "Contractor VPN access expires after 30 days.",
                "Failed VPN login attempts are reviewed daily by IT operations.",
            ],
            "pdf",
        ),
        (
            "Device Setup Checklist",
            "Checklist",
            [
                "Laptop hostname format is MY-CEN-<asset tag>.",
                "Endpoint protection must show healthy status before device handover.",
                "BitLocker recovery key escrow must be confirmed in the device inventory.",
                "New joiners receive Teams, Outlook, VPN, password manager, and browser profile setup.",
            ],
            "docx",
        ),
        (
            "Software Request Standard",
            "Standard",
            [
                "Unlicensed software is not allowed on company devices.",
                "Open-source packages for production work require security review when rated high risk.",
                "Admin rights are temporary and must be approved for a named task.",
                "Software exceptions are reviewed monthly.",
            ],
            "pdf",
        ),
        (
            "Email Security Guide",
            "Guide",
            [
                "Suspected phishing emails must be reported using the Report Phish button.",
                "Users must not forward suspicious attachments to colleagues.",
                "IT security reviews reported emails within one business day.",
                "Credential prompts from unknown links must be treated as high risk.",
            ],
            "docx",
        ),
        (
            "Knowledge Base Admin SOP",
            "SOP",
            [
                "Only System Admin can perform full index rebuilds.",
                "Project Managers can upload and review documents for their own department.",
                "General Employees can search approved content but cannot stage or approve files.",
                "Pending index documents must be updated before chat uses the latest content.",
            ],
            "pdf",
        ),
        (
            "Azure Cost Safety Notes",
            "Notes",
            [
                "Use Azure AI Search Free tier for the student demonstration.",
                "Do not deploy paid Standard search tier without approval.",
                "Blob Storage should only contain small demo documents.",
                "Ollama remains the fallback when Azure OpenAI is unavailable.",
            ],
            "docx",
        ),
        (
            "Incident Ticket Triage",
            "Procedure",
            [
                "Priority 1 tickets require immediate acknowledgement within 15 minutes.",
                "Priority 2 tickets require response within 2 business hours.",
                "Repeated login failures from a foreign location must be escalated to Security.",
                "Resolved tickets must include root cause and user impact summary.",
            ],
            "pdf",
        ),
    ],
    "HR": [
        (
            "Leave Policy 2026",
            "Policy",
            [
                "Annual leave requests should be submitted at least 7 calendar days in advance.",
                "Medical leave requires a valid medical certificate for absences longer than one day.",
                "Emergency leave must be reported to the line manager as soon as practical.",
                "Unused annual leave carry-forward is capped at 5 days unless HR approves an exception.",
            ],
            "docx",
        ),
        (
            "Employee Onboarding Handbook",
            "Handbook",
            [
                "Day 1 onboarding includes ID verification, laptop collection, email setup, and HR briefing.",
                "New employees must complete security awareness training within 5 working days.",
                "Probation review is scheduled before the end of month 3.",
                "The temporary password in the welcome email must be changed immediately.",
            ],
            "pdf",
        ),
        (
            "Expense Claim Guide",
            "Guide",
            [
                "Expense claims must be submitted by the 5th working day of the following month.",
                "Original receipts are required for claims above RM50.",
                "Meal reimbursement during approved travel is capped at RM45 per meal.",
                "Claims without project code are returned for correction.",
            ],
            "docx",
        ),
        (
            "Performance Review Cycle",
            "Procedure",
            [
                "Mid-year check-ins are held in June.",
                "Year-end review forms are due by 15 December.",
                "Managers must record at least two measurable goals for each employee.",
                "Calibration outcomes are visible only to HR and authorized managers.",
            ],
            "pdf",
        ),
        (
            "Remote Work Arrangement",
            "Policy",
            [
                "Hybrid staff must be onsite at least two days per week unless approved otherwise.",
                "Remote work outside Malaysia requires HR and Legal review.",
                "Employees must keep company devices encrypted when working remotely.",
                "Managers review remote-work arrangements every quarter.",
            ],
            "docx",
        ),
        (
            "Training Reimbursement SOP",
            "SOP",
            [
                "Training reimbursement requires manager approval before enrolment.",
                "Certification exam fees are reimbursed only after proof of completion.",
                "Bond terms apply when company-sponsored training exceeds RM3,000.",
                "Training records are stored in the HR learning tracker.",
            ],
            "pdf",
        ),
        (
            "Disciplinary Process Guide",
            "Guide",
            [
                "Misconduct cases must be documented with dates, witnesses, and evidence.",
                "HR must be consulted before issuing a written warning.",
                "Serious misconduct may require immediate suspension pending investigation.",
                "Final disciplinary outcomes are restricted to HR and approved management.",
            ],
            "docx",
        ),
        (
            "Benefits FAQ",
            "FAQ",
            [
                "Dental claims are capped at RM500 per calendar year.",
                "Optical claims are capped at RM300 per calendar year.",
                "Hospitalization benefits follow the active insurance panel.",
                "Benefit eligibility begins after employee profile activation in HRIS.",
            ],
            "pdf",
        ),
    ],
    "ENGINEERING": [
        (
            "Coding Standards",
            "Standard",
            [
                "Pull requests must include tests for behavior changes.",
                "Production branches require at least one reviewer approval.",
                "Secrets must never be committed to source control.",
                "Python modules should keep side effects outside import time where practical.",
            ],
            "docx",
        ),
        (
            "Release Management SOP",
            "SOP",
            [
                "Release candidates are cut every Thursday before 3 PM Malaysia time.",
                "Rollback plans are mandatory for production releases.",
                "Hotfix releases require engineering lead and operations approval.",
                "Release notes must include risk, owner, and verification summary.",
            ],
            "pdf",
        ),
        (
            "API Design Guide",
            "Guide",
            [
                "Public endpoints must return structured error details.",
                "Pagination is required for list endpoints returning more than 100 records.",
                "Breaking API changes require versioning and migration notice.",
                "Backend authorization must not rely on frontend-only checks.",
            ],
            "docx",
        ),
        (
            "Model Evaluation Notes",
            "Notes",
            [
                "Top-K accuracy is measured against labelled expected sources.",
                "Query logs alone cannot prove retrieval accuracy without ground truth.",
                "Miss review logs should separate not-found, permission block, and low-score cases.",
                "Threshold choices must be justified using labelled evaluation results.",
            ],
            "pdf",
        ),
        (
            "Data Pipeline Runbook",
            "Runbook",
            [
                "ETL jobs must log document count, chunk count, and elapsed time.",
                "Failed parsing should not block unrelated documents.",
                "Archived documents are excluded from active retrieval.",
                "Reindex jobs must preserve the existing index until replacement succeeds.",
            ],
            "docx",
        ),
        (
            "Secure Code Review Checklist",
            "Checklist",
            [
                "Check authentication, authorization, input validation, and secret handling.",
                "Review logging for sensitive data exposure.",
                "Confirm dependency changes are necessary and maintained.",
                "Security findings must include impact, evidence, and recommended fix.",
            ],
            "pdf",
        ),
        (
            "Frontend UI Quality Standard",
            "Standard",
            [
                "Operational tools should prioritize dense but readable information.",
                "Buttons should use clear action labels and avoid duplicate meanings.",
                "Long-running actions should use backend jobs where possible.",
                "Tables should default to clean views with details behind expanders.",
            ],
            "docx",
        ),
        (
            "RAG Architecture Decision Record",
            "ADR",
            [
                "Streamlit and Teams are separate clients connected to the same FastAPI backend.",
                "SQLite remains the source of truth for metadata, ACL, audit, and versioning.",
                "Vector stores contain searchable chunks and duplicated filter metadata.",
                "The LLM receives only authorized retrieved context.",
            ],
            "pdf",
        ),
    ],
    "SECURITY": [
        (
            "Access Control Policy",
            "Policy",
            [
                "Access is granted using least privilege and reviewed quarterly.",
                "Privileged access requires named owner, business justification, and expiry date.",
                "Shared accounts are prohibited unless formally approved as a service account.",
                "Access reviews are retained for audit evidence.",
            ],
            "docx",
        ),
        (
            "Incident Response Playbook",
            "Playbook",
            [
                "Security incidents are classified as Low, Medium, High, or Critical.",
                "Critical incidents require notification to Security within 15 minutes.",
                "Containment actions must be recorded before eradication begins.",
                "Post-incident review must list timeline, root cause, and prevention actions.",
            ],
            "pdf",
        ),
        (
            "Data Classification Standard",
            "Standard",
            [
                "Public, Internal, Confidential, and Restricted are the approved data classes.",
                "Restricted data must not be uploaded to unmanaged personal storage.",
                "Confidential files require approved sharing channels.",
                "Classification labels should appear in document headers where practical.",
            ],
            "docx",
        ),
        (
            "Phishing Simulation Procedure",
            "Procedure",
            [
                "Phishing simulations run quarterly.",
                "Users who click simulation links receive targeted micro-training.",
                "Repeat failures are reported to department managers.",
                "Simulation results are reported as aggregate metrics.",
            ],
            "pdf",
        ),
        (
            "Vendor Security Review",
            "Checklist",
            [
                "New vendors handling confidential data require security assessment.",
                "SOC 2, ISO 27001, or equivalent evidence should be requested when relevant.",
                "High-risk findings require mitigation before go-live.",
                "Vendor access must be removed when the contract ends.",
            ],
            "docx",
        ),
        (
            "Key Management Notes",
            "Notes",
            [
                "Production secrets must be stored in approved secret-management systems.",
                "API keys must not appear in screenshots, logs, or demo recordings.",
                "Key rotation is required after suspected exposure.",
                "Demo keys should use minimum required permissions.",
            ],
            "pdf",
        ),
        (
            "Physical Security Guide",
            "Guide",
            [
                "Visitors must wear badges and be escorted in restricted areas.",
                "Tailgating through secured doors must be challenged politely.",
                "Lost access cards must be reported immediately.",
                "Server-room access is limited to authorized IT and facilities staff.",
            ],
            "docx",
        ),
        (
            "Security Awareness FAQ",
            "FAQ",
            [
                "Employees should lock screens when leaving their desks.",
                "Personal email must not be used for company file transfer.",
                "Suspicious USB devices should be handed to IT Security.",
                "Security questions can be sent through the helpdesk queue.",
            ],
            "pdf",
        ),
    ],
    "OPERATIONS": [
        (
            "Service Escalation SOP",
            "SOP",
            [
                "Priority 1 service issues require stakeholder update every 30 minutes.",
                "Priority 2 issues require update every 4 business hours.",
                "Escalation owners must record impact, workaround, and next action.",
                "Customer-facing updates must be approved by the duty manager.",
            ],
            "docx",
        ),
        (
            "Business Continuity Plan",
            "Plan",
            [
                "Critical operations must resume within 4 hours during a declared outage.",
                "Department contact trees are tested twice per year.",
                "Manual workarounds are reviewed after every continuity exercise.",
                "Recovery evidence must be stored in the operations audit folder.",
            ],
            "pdf",
        ),
        (
            "Shift Handover Checklist",
            "Checklist",
            [
                "Outgoing shift must summarize open incidents, risks, and pending approvals.",
                "Incoming shift confirms ownership of unresolved tickets.",
                "High-priority customer commitments are highlighted in the handover note.",
                "Incomplete handovers are escalated to the shift lead.",
            ],
            "docx",
        ),
        (
            "Customer SLA Reference",
            "Reference",
            [
                "Gold SLA response target is 30 minutes.",
                "Silver SLA response target is 2 business hours.",
                "Bronze SLA response target is next business day.",
                "SLA pauses require documented customer dependency.",
            ],
            "pdf",
        ),
        (
            "Change Advisory Board Notes",
            "Notes",
            [
                "Standard changes may be pre-approved when risk is low and repeatable.",
                "Emergency changes require post-implementation review.",
                "CAB minutes must list decision, risk, owner, and implementation window.",
                "Rejected changes should include a clear reason and resubmission path.",
            ],
            "docx",
        ),
        (
            "Knowledge Article Quality Guide",
            "Guide",
            [
                "Knowledge articles should include symptoms, cause, resolution, and validation.",
                "Article owners review high-use articles every 90 days.",
                "Screenshots must not expose customer personal data.",
                "Outdated articles should be archived instead of left active.",
            ],
            "pdf",
        ),
        (
            "Operational Risk Register",
            "Register",
            [
                "Risks are scored by likelihood and impact.",
                "High risks require named mitigation owner and due date.",
                "Residual risk is reviewed after mitigation is complete.",
                "Risk register updates are reviewed monthly.",
            ],
            "docx",
        ),
        (
            "Facilities Request Guide",
            "Guide",
            [
                "Facilities requests should include location, urgency, and affected users.",
                "Air-conditioning issues in server rooms are treated as urgent.",
                "Access-card requests require manager approval.",
                "Completed facilities tasks are closed with date and action taken.",
            ],
            "pdf",
        ),
    ],
}


ONENOTE_PAGES = {
    "IT": [
        ("Password Reset Quick Notes", "Quick Notes", "Password resets use self-service first. Helpdesk resets require identity verification using employee ID and manager name."),
        ("VPN Troubleshooting Cases", "Runbook Notes", "Common VPN fixes: check MFA prompt, confirm device compliance, refresh profile, then escalate with logs."),
        ("Index Rebuild Demo Notes", "RAG Admin", "Full rebuild is admin-only. Incremental update is preferred when only one approved document changed."),
        ("Graph Connector Setup Notes", "Connector Notes", "OneDrive root path is Enterprise Knowledge Base. Department is inferred from the first folder under the root."),
    ],
    "HR": [
        ("Leave Approval Examples", "Policy Notes", "Leave should be requested 7 days early. Emergency leave can be reported immediately and documented later."),
        ("Onboarding Day 1 Notes", "Onboarding", "Day 1 includes HR briefing, laptop pickup, account activation, and mandatory security training assignment."),
        ("Benefits Claim Reminders", "Benefits", "Dental cap is RM500 yearly. Optical cap is RM300 yearly. Receipts are required for reimbursement."),
        ("Misconduct Review Notes", "Employee Relations", "Misconduct records must include evidence, timeline, witnesses, and HR consultation before final action."),
    ],
    "ENGINEERING": [
        ("Pull Request Review Notes", "Development", "Reviewers check tests, secrets, breaking changes, performance risk, and rollback notes before approval."),
        ("RAG Evaluation Notes", "AI Quality", "Top-K accuracy needs labelled expected sources. Query logs alone are useful signals but not ground truth."),
        ("Release Checklist Notes", "Release", "Every release needs owner, risk, verification, rollback plan, and customer-impact statement."),
        ("API Error Handling Notes", "Backend", "API errors should be structured, actionable, and safe to show in frontend status panels."),
    ],
    "SECURITY": [
        ("Incident Severity Notes", "Incident Response", "Critical incidents require Security notification within 15 minutes and recorded containment steps."),
        ("Data Classification Examples", "Governance", "Restricted data cannot be uploaded to unmanaged personal storage. Confidential data requires approved sharing."),
        ("Phishing Drill Notes", "Awareness", "Simulation failures trigger micro-training. Repeat failures are reported as aggregate manager metrics."),
        ("Secret Handling Notes", "Engineering Security", "Never expose API keys in screenshots, logs, commits, or demo recordings."),
    ],
    "OPERATIONS": [
        ("Shift Handover Notes", "Service Desk", "Handover notes include open incidents, customer commitments, risks, and pending approvals."),
        ("SLA Escalation Examples", "Customer Ops", "Gold response target is 30 minutes. Silver is 2 business hours. Bronze is next business day."),
        ("Continuity Exercise Notes", "BCP", "Critical operations target recovery within 4 hours during a declared outage."),
        ("Knowledge Article Notes", "Knowledge Ops", "High-use articles are reviewed every 90 days and outdated articles should be archived."),
    ],
}


def slugify(value: str) -> str:
    """Convert a title into a filesystem-friendly lowercase slug."""
    return re.sub(r"[^a-z0-9]+", "_", value.lower()).strip("_")


def wrap_lines(lines: list[str], width: int = 92) -> list[str]:
    """Wrap document content so generated PDFs and Markdown stay readable."""
    wrapped: list[str] = []
    for line in lines:
        wrapped.extend(textwrap.wrap(line, width=width) or [""])
    return wrapped


def build_document_lines(department: str, title: str, category: str, points: list[str]) -> list[str]:
    """Create consistent synthetic policy text with metadata-like cues."""
    lines = [
        title,
        f"Department: {department}",
        f"Category: {category}",
        "Source: Synthetic Enterprise Knowledge Base",
        "Audience: Internal staff with approved department access",
        "",
        "Purpose",
        f"This document provides department-specific guidance for {department} staff and approved collaborators.",
        "",
        "Key Requirements",
    ]
    lines.extend(f"- {point}" for point in points)
    lines.extend(
        [
            "",
            "Review And Ownership",
            f"The {department} document owner reviews this content every 90 days or after a material process change.",
            "Outdated versions should be archived before the replacement version is indexed.",
            "",
            "RAG Test Notes",
            f"Queries about '{title.lower()}' should cite this source when the user has access to {department}.",
            "If a user from another department lacks permission, the system should not use this content in the answer.",
        ]
    )
    return lines


def save_docx(path: Path, title: str, lines: list[str]) -> None:
    """Write a DOCX file using the installed python-docx package."""
    document = Document()
    document.add_heading(title, level=1)
    for line in lines[1:]:
        if not line:
            document.add_paragraph()
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line in {"Purpose", "Key Requirements", "Review And Ownership", "RAG Test Notes"}:
            document.add_heading(line, level=2)
        else:
            document.add_paragraph(line)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def pdf_escape(text: str) -> str:
    """Escape text for a simple PDF text stream."""
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def save_pdf(path: Path, title: str, lines: list[str]) -> None:
    """Write a small extractable-text PDF without adding extra dependencies."""
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf_lines = wrap_lines(lines)
    text_ops = ["BT", "/F1 11 Tf", "50 780 Td", "14 TL"]
    for index, line in enumerate(pdf_lines[:48]):
        if index == 0:
            text_ops.append(f"({pdf_escape(line)}) Tj")
        else:
            text_ops.append("T*")
            text_ops.append(f"({pdf_escape(line)}) Tj")
    text_ops.append("ET")
    stream = "\n".join(text_ops).encode("latin-1", "replace")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]

    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(content))
        content.extend(f"{number} 0 obj\n".encode("ascii"))
        content.extend(obj)
        content.extend(b"\nendobj\n")
    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(content)


def save_onenote_page(path: Path, department: str, title: str, section: str, body: str) -> None:
    """Write a Markdown page that can be pasted into OneNote as page content."""
    path.parent.mkdir(parents=True, exist_ok=True)
    content = "\n".join(
        [
            f"# {title}",
            "",
            f"Department: {department}",
            f"Notebook: Enterprise Knowledge Base",
            f"Section: {section}",
            "Source: Synthetic OneNote page",
            "",
            body,
            "",
            "Access note: this page should be searchable only for users allowed to retrieve this department's content.",
        ]
    )
    path.write_text(content + "\n", encoding="utf-8")


def main() -> None:
    """Generate OneDrive files, OneNote pages, and a manifest."""
    manifest_rows: list[dict[str, str]] = []

    for department, docs in DEPARTMENTS.items():
        for title, category, points, extension in docs:
            lines = build_document_lines(department, title, category, points)
            filename = f"{slugify(title)}.{extension}"
            path = ONEDRIVE_ROOT / department / filename
            if extension == "docx":
                save_docx(path, title, lines)
            else:
                save_pdf(path, title, lines)
            manifest_rows.append(
                {
                    "platform": "OneDrive",
                    "department": department,
                    "title": title,
                    "category": category,
                    "file_type": extension.upper(),
                    "relative_path": str(path.relative_to(ROOT)),
                }
            )

    for department, pages in ONENOTE_PAGES.items():
        for title, section, body in pages:
            path = ONENOTE_ROOT / department / section / f"{slugify(title)}.md"
            save_onenote_page(path, department, title, section, body)
            manifest_rows.append(
                {
                    "platform": "OneNote",
                    "department": department,
                    "title": title,
                    "category": section,
                    "file_type": "MD",
                    "relative_path": str(path.relative_to(ROOT)),
                }
            )

    ROOT.mkdir(parents=True, exist_ok=True)
    with (ROOT / "manifest.csv").open("w", newline="", encoding="utf-8") as csv_file:
        writer = csv.DictWriter(
            csv_file,
            fieldnames=["platform", "department", "title", "category", "file_type", "relative_path"],
        )
        writer.writeheader()
        writer.writerows(manifest_rows)

    readme = "\n".join(
        [
            "# Enterprise Knowledge Base Synthetic Dataset",
            "",
            "Generated for Searchable RAG Copilot connector testing.",
            "",
            "## Structure",
            "",
            "- `onedrive/Enterprise Knowledge Base/<DEPARTMENT>/...` contains DOCX/PDF files for OneDrive upload.",
            "- `onenote/Enterprise Knowledge Base/<DEPARTMENT>/<SECTION>/...` contains Markdown pages that can be pasted into OneNote pages.",
            "- `manifest.csv` lists every generated item.",
            "",
            "## Departments",
            "",
            "- IT",
            "- HR",
            "- ENGINEERING",
            "- SECURITY",
            "- OPERATIONS",
            "",
            "## Intended Connector Behavior",
            "",
            "The first folder or OneNote section path after `Enterprise Knowledge Base` should infer the department.",
            "All documents are synthetic and safe for demo use.",
            "",
            f"Generated items: {len(manifest_rows)}",
        ]
    )
    (ROOT / "README.md").write_text(readme + "\n", encoding="utf-8")

    print(f"Generated {len(manifest_rows)} items under {ROOT}")
    print(f"OneDrive files: {sum(row['platform'] == 'OneDrive' for row in manifest_rows)}")
    print(f"OneNote pages: {sum(row['platform'] == 'OneNote' for row in manifest_rows)}")


if __name__ == "__main__":
    main()
